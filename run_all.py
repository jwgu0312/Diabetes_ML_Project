"""Execute all notebooks and build final Markdown/DOCX reports."""

from __future__ import annotations

import sys
from pathlib import Path

import nbformat
import pandas as pd
from nbclient import NotebookClient

from project_utils import FIGURES_DIR, REPORTS_DIR, ROOT, TABLES_DIR


NOTEBOOKS = [
    "01_Data_Inspection_and_EDA.ipynb",
    "02_Data_Preprocessing.ipynb",
    "03_Baseline_Models.ipynb",
    "04_Advanced_Models_and_Tuning.ipynb",
    "05_Ensemble_Learning.ipynb",
    "06_Model_Explainability_SHAP.ipynb",
    "07_Final_Evaluation_and_Comparison.ipynb",
]


def execute_notebooks() -> None:
    for filename in NOTEBOOKS:
        path = ROOT / "code" / filename
        print(f"\n=== Executing {filename} ===")
        with path.open("r", encoding="utf-8") as f:
            nb = nbformat.read(f, as_version=4)
        client = NotebookClient(
            nb,
            timeout=3600,
            kernel_name="python3",
            resources={"metadata": {"path": str(ROOT)}},
            allow_errors=False,
        )
        client.execute()
        with path.open("w", encoding="utf-8") as f:
            nbformat.write(nb, f)
        print(f"Completed {filename}")


def table_md(path: Path, rows: int = 12) -> str:
    if not path.exists():
        return f"Table not found: `{path.name}`"
    df = pd.read_csv(path)
    return df.head(rows).to_markdown(index=False)


def build_markdown_report() -> Path:
    final = pd.read_csv(TABLES_DIR / "final_model_comparison.csv")
    best = final.iloc[0]
    class_dist = pd.read_csv(TABLES_DIR / "class_distribution.csv")
    cleaning = pd.read_csv(TABLES_DIR / "cleaning_summary.csv")
    smote = pd.read_csv(TABLES_DIR / "smote_class_distribution.csv")
    shap_path = TABLES_DIR / "shap_feature_importance.csv"
    shap_table = pd.read_csv(shap_path) if shap_path.exists() else pd.DataFrame()
    top_features = ", ".join(shap_table["Feature"].head(8).tolist()) if not shap_table.empty else "not available"

    content = f"""# Final Report: Multi-Class Diabetes Classification Using BRFSS 2015

## Executive Summary

This project developed a complete, reproducible machine learning workflow for predicting diabetes status from the BRFSS 2015 health indicators dataset. The target variable was `Diabetes_012`, a three-class outcome: no diabetes, prediabetes, and diabetes. The workflow included data inspection, cleaning, exploratory data analysis, preprocessing, SMOTE class balancing, six supervised machine learning models, hyperparameter tuning, stacking ensemble learning, SHAP explainability, publication-quality visualizations, and final model comparison.

The best-ranked model in the final comparison was **{best['Model']}**, with macro F1-score **{best['F1_macro']:.3f}**, accuracy **{best['Accuracy']:.3f}**, and macro ROC-AUC **{best['ROC_AUC_ovr_macro']:.3f}**. Because the dataset is imbalanced, macro F1-score is emphasized because it gives equal importance to no diabetes, prediabetes, and diabetes.

## Dataset and Problem Definition

The dataset contains self-reported BRFSS 2015 health indicators. Each row represents a survey participant and each column represents a health, demographic, or access-to-care variable. The modeling task is multi-class classification:

- `0`: no diabetes
- `1`: prediabetes
- `2`: diabetes

The project is framed as a public health prediction problem. The goal is not to replace clinical diagnosis, but to demonstrate how machine learning can identify patterns associated with diabetes risk using population-level survey indicators.

## Data Inspection and Cleaning

Initial inspection confirmed the dataset dimensions, missing-value status, duplicate records, and class balance. The dataset contained no missing cells, so imputation was not required. Exact duplicate records were removed before modeling to reduce repeated-observation bias.

Cleaning summary:

{table_md(TABLES_DIR / 'cleaning_summary.csv')}

Class distribution:

{table_md(TABLES_DIR / 'class_distribution.csv')}

The class distribution shows substantial imbalance. The no-diabetes class is much larger than the prediabetes class. This matters because a model can achieve high accuracy by focusing on the majority class while performing poorly on prediabetes.

## Exploratory Data Analysis

The EDA stage generated descriptive statistics, class distribution plots, BMI distribution plots, BMI-by-class boxplots, a full correlation heatmap, and selected feature means by diabetes status. These figures are saved as both PNG and PDF files in `results/figures`.

Key EDA findings:

- Diabetes status is imbalanced, with prediabetes being the rarest class.
- BMI tends to be higher among participants with diabetes or prediabetes than among participants without diabetes.
- General health, age group, high blood pressure, high cholesterol, difficulty walking, and physical health indicators show meaningful variation across diabetes classes.
- Socioeconomic indicators such as income and education provide useful context for public health interpretation.

Important figures:

- `eda_class_distribution.png`
- `eda_bmi_distribution.png`
- `eda_bmi_by_diabetes_status.png`
- `eda_correlation_matrix.png`
- `eda_feature_means_by_class.png`

## Preprocessing

The data was split into training and test sets using stratified sampling. Stratification preserves the class proportions in both sets, which is essential for imbalanced multi-class classification.

All predictor variables were scaled using `StandardScaler`. Scaling is especially important for Logistic Regression and Support Vector Machine models because these algorithms are sensitive to feature magnitude.

SMOTE was applied only to the training data. This avoids test-set leakage and ensures that evaluation remains realistic.

SMOTE before/after distribution:

{table_md(TABLES_DIR / 'smote_class_distribution.csv')}

## Why SMOTE Is Important

SMOTE, or Synthetic Minority Oversampling Technique, creates synthetic training examples for minority classes. In this project, prediabetes is rare relative to the no-diabetes class. Without imbalance handling, models may learn decision boundaries that under-detect prediabetes. SMOTE helps the model see a more balanced training signal, improving the chance that minority-class patterns are learned.

SMOTE does not change the test set. The held-out test set remains in the original distribution, which keeps final evaluation honest.

## Models Trained

The project trained the six required machine learning models:

- Logistic Regression
- Decision Tree
- Random Forest
- XGBoost
- LightGBM
- Support Vector Machine

The project also trained a `StackingClassifier` ensemble that combines Random Forest, XGBoost, and LightGBM base learners with Logistic Regression as the meta-model.

## Baseline Model Results

Baseline models were evaluated using accuracy, macro precision, macro recall, macro F1-score, weighted F1-score, macro ROC-AUC, and macro PR-AUC.

{table_md(TABLES_DIR / 'model_performance.csv')}

Accuracy is reported but not treated as the only success measure. Macro F1-score is more informative in this setting because each class contributes equally to the metric.

## Hyperparameter Tuning

Random Forest, XGBoost, and LightGBM were tuned with `RandomizedSearchCV`. The tuning objective was macro F1-score, which supports balanced performance across all three classes.

Tuned model results:

{table_md(TABLES_DIR / 'tuned_model_performance.csv')}

Hyperparameter tuning helps control model complexity. For tree ensembles, settings such as tree depth, number of estimators, learning rate, and number of leaves influence the trade-off between underfitting and overfitting.

## Ensemble Learning

The stacking ensemble uses multiple base models and learns how to combine their probability outputs. Ensemble learning can help because different algorithms capture different structures:

- Random Forest reduces variance through bagging.
- XGBoost learns additive boosted trees that focus on difficult cases.
- LightGBM provides efficient gradient boosting with strong nonlinear modeling capacity.
- Logistic Regression as the meta-model learns how much to trust each base learner for the final prediction.

Stacking result:

{table_md(TABLES_DIR / 'stacking_model_performance.csv')}

## Final Model Comparison

The final comparison ranks baseline, tuned, and ensemble models using macro F1-score and ROC-AUC.

{table_md(TABLES_DIR / 'final_model_comparison.csv', rows=20)}

The final ranking should be interpreted in the context of class imbalance. A model with slightly lower accuracy but better macro recall or macro F1-score may be more appropriate for public health screening because it gives more attention to minority disease categories.

## Explainable AI with SHAP

SHAP was used to explain the tree-based diabetes prediction model. SHAP values estimate how much each feature contributes to model predictions. The most important features in this project were:

**{top_features}**

Top SHAP features:

{table_md(TABLES_DIR / 'shap_feature_importance.csv')}

Public health interpretation:

- BMI reflects body composition and metabolic risk.
- General health is a broad self-reported indicator of overall disease burden.
- High blood pressure and high cholesterol are clinically meaningful cardiometabolic risk factors.
- Age captures increasing chronic disease risk across the lifespan.
- Difficulty walking and physical health may reflect mobility limitations and comorbidity.
- Income and education can reflect social determinants of health and access to preventive care.

SHAP explanations are model explanations, not causal proof. They show which variables the trained model uses most strongly when making predictions.

## Figures Generated

All figures were saved in both PNG and PDF format at 300 dpi. This supports both presentation use and academic submission. Major figures include:

- Class distribution before and after SMOTE
- BMI distribution and BMI by diabetes status
- Correlation heatmap
- Model comparison charts
- Confusion matrices
- ROC curves
- Precision-recall curves
- SHAP summary plot
- SHAP feature-importance plot

## Tables Generated

All major tables were saved as both CSV and XLSX files in `results/tables`, including:

- Descriptive statistics
- Missing-value analysis
- Duplicate analysis
- Class distribution
- Train/test split summary
- SMOTE class distribution
- Classification reports
- Hyperparameter search results
- Feature-importance rankings
- Final model comparison
- Artifact verification

## Reproducibility

Reproducibility was supported through:

- Fixed random seed: `42`
- Relative paths from the project root
- Modular utility functions in `code/project_utils.py`
- Notebook-by-notebook execution
- Saved models using `joblib`
- Saved tables as CSV and XLSX
- Saved figures as PNG and PDF
- `requirements.txt` with required packages

To reproduce the project from the project root:

```bash
python -m pip install -r requirements.txt
python code/generate_project.py
python code/run_all_and_build_report.py
```

## Presentation-Ready Speaking Points

1. This project predicts three diabetes-status categories using the BRFSS 2015 health indicators dataset.
2. The dataset is highly imbalanced, so macro F1-score is more meaningful than accuracy alone.
3. Duplicate rows were removed and no missing-value imputation was necessary.
4. Features were standardized to support scale-sensitive algorithms such as Logistic Regression and SVM.
5. SMOTE was applied only to the training set to improve minority-class learning without contaminating the test set.
6. Six required models were trained and evaluated with a consistent held-out test set.
7. Random Forest, XGBoost, and LightGBM were tuned with cross-validation.
8. A stacking ensemble was trained to combine complementary model strengths.
9. SHAP was used to explain which health indicators were most influential in model predictions.
10. The strongest predictors are interpretable public health variables, including BMI, general health, blood pressure, cholesterol, age, and mobility-related health.

## Limitations

The BRFSS dataset is observational and self-reported, so model results should not be interpreted as causal medical evidence. The prediabetes class is small, making it challenging to classify. The dataset does not include all clinical measurements that physicians would use for diagnosis, such as laboratory glucose or HbA1c values.

## Conclusion

This project demonstrates a complete academic machine learning workflow for a real public health classification problem. It combines rigorous preprocessing, class imbalance handling, multiple model families, hyperparameter tuning, ensemble learning, explainable AI, and reproducible reporting. The final outputs are suitable for university submission, team presentation, and further extension.
"""
    out = REPORTS_DIR / "Final_Report.md"
    out.write_text(content, encoding="utf-8")
    return out


def build_docx_report(markdown_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:
        raise ImportError("python-docx is required to create Final_Report.docx") from exc

    doc = Document()
    doc.add_heading("Final Report: Multi-Class Diabetes Classification Using BRFSS 2015", 0)
    text = markdown_path.read_text(encoding="utf-8")
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("# "):
            continue
        if block.startswith("## "):
            doc.add_heading(block.replace("## ", ""), level=1)
        elif block.startswith("### "):
            doc.add_heading(block.replace("### ", ""), level=2)
        elif block.startswith("- "):
            for line in block.splitlines():
                if line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
        elif block.startswith("1. "):
            for line in block.splitlines():
                if ". " in line:
                    doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
        elif block.startswith("|"):
            rows = [line for line in block.splitlines() if line.startswith("|") and "---" not in line]
            if rows:
                parsed = [[cell.strip() for cell in row.strip("|").split("|")] for row in rows]
                table = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
                table.style = "Table Grid"
                for i, row in enumerate(parsed):
                    for j, value in enumerate(row):
                        table.cell(i, j).text = value
        elif block.startswith("```"):
            doc.add_paragraph(block.replace("```bash", "").replace("```", "").strip())
        else:
            doc.add_paragraph(block)

    for fig_name in [
        "eda_class_distribution.png",
        "final_model_comparison_metrics.png",
        "shap_feature_importance_plot.png",
        "shap_summary_plot.png",
    ]:
        fig_path = FIGURES_DIR / fig_name
        if fig_path.exists():
            doc.add_heading(fig_name.replace("_", " ").replace(".png", "").title(), level=1)
            doc.add_picture(str(fig_path), width=Inches(6.2))

    out = REPORTS_DIR / "Final_Report.docx"
    doc.save(out)
    return out


def verify_final_artifacts() -> None:
    expected = [
        ROOT / "requirements.txt",
        REPORTS_DIR / "Final_Report.md",
        REPORTS_DIR / "Final_Report.docx",
        TABLES_DIR / "final_model_comparison.csv",
        TABLES_DIR / "final_model_comparison.xlsx",
        TABLES_DIR / "shap_feature_importance.csv",
        FIGURES_DIR / "shap_summary_plot.png",
        FIGURES_DIR / "shap_summary_plot.pdf",
        ROOT / "results" / "models" / "stacking_ensemble.joblib",
    ]
    missing = [str(p.relative_to(ROOT)) for p in expected if not p.exists()]
    if missing:
        raise FileNotFoundError(f"Missing final artifacts: {missing}")
    fig_png = list(FIGURES_DIR.glob("*.png"))
    table_csv = list(TABLES_DIR.glob("*.csv"))
    print(f"Verified {len(fig_png)} PNG figures and {len(table_csv)} CSV tables.")


def main():
    execute_notebooks()
    md_path = build_markdown_report()
    docx_path = build_docx_report(md_path)
    verify_final_artifacts()
    print(f"Markdown report: {md_path}")
    print(f"DOCX report: {docx_path}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise
